from __future__ import annotations

import csv
import io
import mimetypes
import re
import zipfile
from contextlib import suppress
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document

from synextra.services.block_chunker import ChunkedText, chunk_pdf_blocks, chunk_text_pages
from synextra.services.document_store import PageText, build_page_texts_from_blocks
from synextra.services.pdf_ingestion import (
    PdfEncryptedError,
    PdfIngestionError,
    extract_pdf_blocks,
    sha256_hex,
)


class DocumentIngestionError(RuntimeError):
    """Base exception for ingestion of non-PDF documents."""


class UnsupportedDocumentTypeError(DocumentIngestionError):
    """Raised when an uploaded document type is not supported."""


class DocumentParseError(DocumentIngestionError):
    """Raised when a supported document cannot be parsed."""


DocumentKind = Literal["pdf", "doc", "docx", "csv", "xlsx", "text"]


_OLE_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


_EXT_TO_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".csv": "text/csv",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


_CODE_EXTENSIONS: set[str] = {
    ".py",
    ".pyi",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".xml",
    ".html",
    ".css",
    ".scss",
    ".sql",
    ".sh",
    ".bash",
    ".ps1",
    ".go",
    ".rs",
    ".java",
    ".kt",
    ".swift",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
    ".cc",
    ".cs",
    ".rb",
    ".php",
}


def _is_pdf_bytes(
    data: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> bool:
    if content_type and content_type.lower().startswith("application/pdf"):
        return True
    if filename and filename.lower().endswith(".pdf"):
        return True
    return data.startswith(b"%PDF")


def _is_probably_text(data: bytes) -> bool:
    if not data:
        return True
    head = data[:4096]
    if head.count(b"\x00") / max(1, len(head)) > 0.02:
        return False
    try:
        head.decode("utf-8")
        return True
    except UnicodeDecodeError:
        # ASCII-ish heuristic.
        printable = sum(1 for b in head if 9 <= b <= 13 or 32 <= b <= 126)
        return printable / max(1, len(head)) > 0.9


def _detect_zip_kind(data: bytes) -> DocumentKind | None:
    if not data.startswith(b"PK"):
        return None
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = set(zf.namelist())
    except Exception:
        return None

    # Office Open XML formats are ZIP containers with well-known internal paths.
    if "word/document.xml" in names:
        return "docx"
    if "xl/workbook.xml" in names:
        return "xlsx"
    return None


def detect_document_kind(
    *,
    data: bytes,
    filename: str | None,
    content_type: str | None,
) -> tuple[DocumentKind, str]:
    """Detect the document kind and a best-effort mime type."""

    if _is_pdf_bytes(data, filename=filename, content_type=content_type):
        return "pdf", "application/pdf"

    ext = Path(filename).suffix.lower() if filename else ""

    # Prefer explicit extensions.
    if ext in {".docx", ".xlsx"}:
        kind: DocumentKind = "docx" if ext == ".docx" else "xlsx"
        return kind, _EXT_TO_MIME.get(ext, "application/octet-stream")
    if ext == ".doc":
        return "doc", _EXT_TO_MIME[".doc"]
    if ext == ".csv":
        return "csv", _EXT_TO_MIME[".csv"]
    if ext in {".txt", ".md"}:
        return "text", _EXT_TO_MIME.get(ext, "text/plain")
    if ext in _CODE_EXTENSIONS:
        # Treat code/config as plain text.
        guessed = mimetypes.guess_type(filename or "")[0] or "text/plain"
        return "text", guessed

    # Sniff OOXML ZIP containers.
    zip_kind = _detect_zip_kind(data)
    if zip_kind is not None:
        if zip_kind == "docx":
            return "docx", _EXT_TO_MIME[".docx"]
        if zip_kind == "xlsx":
            return "xlsx", _EXT_TO_MIME[".xlsx"]

    # Legacy .doc is an OLE container.
    if data.startswith(_OLE_SIGNATURE):
        return "doc", _EXT_TO_MIME[".doc"]

    # Content-type hints.
    if content_type:
        lowered = content_type.lower()
        if lowered.startswith("text/"):
            return "text", content_type
        if lowered in {"application/json", "application/xml"}:
            return "text", content_type
        if lowered in {
            _EXT_TO_MIME[".docx"],
            _EXT_TO_MIME[".xlsx"],
            _EXT_TO_MIME[".doc"],
            _EXT_TO_MIME[".csv"],
        }:
            # If the server gave us an OOXML/doc/csv content type but no extension,
            # assume it is correct.
            if lowered == _EXT_TO_MIME[".docx"]:
                return "docx", _EXT_TO_MIME[".docx"]
            if lowered == _EXT_TO_MIME[".xlsx"]:
                return "xlsx", _EXT_TO_MIME[".xlsx"]
            if lowered == _EXT_TO_MIME[".doc"]:
                return "doc", _EXT_TO_MIME[".doc"]
            if lowered == _EXT_TO_MIME[".csv"]:
                return "csv", _EXT_TO_MIME[".csv"]

    # Fallback: treat as text if it looks like text.
    if _is_probably_text(data):
        return "text", "text/plain"

    supported = ", ".join(sorted(set(_EXT_TO_MIME) | _CODE_EXTENSIONS))
    raise UnsupportedDocumentTypeError(
        f"Unsupported document type. Supported extensions include: {supported}"
    )


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        # Best-effort fallback.
        return data.decode("utf-8", errors="replace")


def _normalize_lines(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    for line in lines:
        # Preserve indentation for code; trim trailing whitespace.
        normalized = str(line).rstrip()
        cleaned.append(normalized)
    # Strip trailing empty lines.
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return cleaned


def _paginate_lines(lines: list[str], *, lines_per_page: int = 160) -> list[PageText]:
    lines = _normalize_lines(lines)
    if not lines:
        return [PageText(page_number=0, lines=[], line_count=0)]

    pages: list[PageText] = []
    for page_number, offset in enumerate(range(0, len(lines), max(1, lines_per_page))):
        page_lines = lines[offset : offset + lines_per_page]
        pages.append(
            PageText(
                page_number=page_number,
                lines=page_lines,
                line_count=len(page_lines),
            )
        )
    return pages


def _extract_docx_lines(data: bytes) -> list[str]:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError("Failed to parse .docx document") from exc

    lines: list[str] = []
    for paragraph in getattr(doc, "paragraphs", []):
        text = " ".join(str(getattr(paragraph, "text", "")).split())
        if text:
            lines.append(text)

    # Tables.
    for table in getattr(doc, "tables", []):
        for row in getattr(table, "rows", []):
            cells = getattr(row, "cells", [])
            values = []
            for cell in cells:
                cell_text = " ".join(str(getattr(cell, "text", "")).split())
                values.append(cell_text)
            # Trim trailing empties.
            while values and not values[-1]:
                values.pop()
            if values:
                lines.append(" | ".join(values))

    return lines


def _extract_csv_lines(data: bytes) -> list[str]:
    raw = _decode_text(data)
    reader = csv.reader(io.StringIO(raw))
    lines: list[str] = []
    for row in reader:
        values = [str(cell).strip() for cell in row]
        while values and not values[-1]:
            values.pop()
        if values:
            lines.append(", ".join(values))
    return lines


def _extract_xlsx_lines(data: bytes) -> list[str]:
    try:
        import openpyxl  # type: ignore[import-untyped]
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError("openpyxl is required to parse .xlsx files") from exc

    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(data),
            data_only=True,
            read_only=True,
        )
    except Exception as exc:
        raise DocumentParseError("Failed to parse .xlsx workbook") from exc

    lines: list[str] = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"# Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                values = []
                for value in row:
                    if value is None:
                        values.append("")
                    else:
                        values.append(str(value))
                while values and not str(values[-1]).strip():
                    values.pop()
                if not values:
                    continue
                lines.append(" | ".join(values))
            lines.append("")
    finally:
        with suppress(Exception):
            wb.close()

    return lines


_RTF_HEADER_RE = re.compile(r"^\{\\rtf", re.IGNORECASE)
_RTF_CONTROL_RE = re.compile(r"\\[a-zA-Z]+-?\d* ?")
_RTF_BRACE_RE = re.compile(r"[{}]")


def _strip_rtf(text: str) -> str:
    # Extremely small/naive RTF stripper.
    text = _RTF_BRACE_RE.sub("", text)
    text = _RTF_CONTROL_RE.sub("", text)
    text = text.replace("\\'", "")
    return " ".join(text.split())


def _extract_doc_lines(data: bytes) -> list[str]:
    # Some .doc files are actually RTF.
    head = data[:64]
    try:
        head_text = head.decode("ascii", errors="ignore")
    except Exception:
        head_text = ""

    if _RTF_HEADER_RE.search(head_text):
        raw = data.decode("latin-1", errors="ignore")
        stripped = _strip_rtf(raw)
        return [line for line in stripped.split("\\n") if line.strip()]

    # Best-effort strings extraction (OLE/legacy Word binary).
    ascii_runs = re.findall(rb"[\x20-\x7E]{4,}", data)
    ascii_lines = [run.decode("ascii", errors="ignore") for run in ascii_runs]

    # UTF-16LE extraction captures many Word documents.
    try:
        utf16_text = data.decode("utf-16le", errors="ignore")
    except Exception:
        utf16_text = ""

    utf16_candidates = re.findall(r"[A-Za-z0-9][A-Za-z0-9\s\.,;:'\"!?()\-_/]{8,}", utf16_text)

    combined = ascii_lines + utf16_candidates
    lines: list[str] = []
    for item in combined:
        cleaned = " ".join(str(item).split())
        if len(cleaned) < 6:
            continue
        # Require at least one letter to avoid pure formatting noise.
        if not re.search(r"[A-Za-z]", cleaned):
            continue
        lines.append(cleaned)

    # De-duplicate while preserving order.
    seen: set[str] = set()
    unique: list[str] = []
    for line in lines:
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(line)

    if not unique:
        raise DocumentParseError("Failed to extract text from .doc file")

    return unique


@dataclass(frozen=True)
class ParsedDocument:
    kind: DocumentKind
    mime_type: str
    checksum_sha256: str
    page_count: int
    pages: list[PageText]
    chunks: list[ChunkedText]


def parse_document(
    *,
    data: bytes,
    filename: str | None,
    content_type: str | None,
    document_id: str | None = None,
    lines_per_page: int = 160,
) -> ParsedDocument:
    """Parse + chunk a supported document into pages and retrieval chunks."""

    resolved_id = document_id or sha256_hex(data)
    checksum = sha256_hex(data)

    kind, mime = detect_document_kind(data=data, filename=filename, content_type=content_type)

    if kind == "pdf":
        try:
            pdf = extract_pdf_blocks(data, sort=True)
        except PdfEncryptedError:
            raise
        except PdfIngestionError as exc:
            raise DocumentParseError("Failed to parse PDF") from exc

        chunks = chunk_pdf_blocks(document_id=resolved_id, blocks=pdf.blocks)
        pages = build_page_texts_from_blocks(pdf.blocks, pdf.page_count)
        return ParsedDocument(
            kind=kind,
            mime_type=mime,
            checksum_sha256=pdf.checksum_sha256,
            page_count=pdf.page_count,
            pages=pages,
            chunks=chunks,
        )

    try:
        if kind == "docx":
            lines = _extract_docx_lines(data)
        elif kind == "doc":
            lines = _extract_doc_lines(data)
        elif kind == "csv":
            lines = _extract_csv_lines(data)
        elif kind == "xlsx":
            lines = _extract_xlsx_lines(data)
        else:
            # text
            lines = _decode_text(data).splitlines()
    except DocumentParseError:
        raise
    except PdfEncryptedError:
        raise
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError("Failed to parse document") from exc

    pages = _paginate_lines(lines, lines_per_page=lines_per_page)
    chunks = chunk_text_pages(document_id=resolved_id, pages=pages)
    return ParsedDocument(
        kind=kind,
        mime_type=mime,
        checksum_sha256=checksum,
        page_count=len(pages),
        pages=pages,
        chunks=chunks,
    )
